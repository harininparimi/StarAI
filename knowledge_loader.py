from pathlib import Path
from docx import Document
from openpyxl import load_workbook
from config import SUPPORTED_KNOWLEDGE_EXTENSIONS


def _read_docx(path: Path) -> str:
    doc = Document(path)
    lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            values = [c.text.strip().replace("\n", " ") for c in row.cells]
            if any(values):
                lines.append(" | ".join(values))
    return "\n".join(lines)


def _read_xlsx(path: Path) -> str:
    wb = load_workbook(path, data_only=True, read_only=True)
    lines = []
    for ws in wb.worksheets:
        lines.append(f"Sheet: {ws.title}")
        for row in ws.iter_rows(values_only=True):
            values = ["" if v is None else str(v) for v in row]
            if any(values):
                lines.append(" | ".join(values))
    return "\n".join(lines)


def load_knowledge(folder: Path) -> tuple[str, list[str]]:
    files = sorted(
        p for p in folder.iterdir()
        if p.is_file()
        and not p.name.startswith("~$")
        and p.suffix.lower() in SUPPORTED_KNOWLEDGE_EXTENSIONS
    )
    if not files:
        raise FileNotFoundError("No StarAI knowledge files were found.")

    sections = []
    for path in files:
        if path.suffix.lower() == ".docx":
            content = _read_docx(path)
        elif path.suffix.lower() == ".xlsx":
            content = _read_xlsx(path)
        else:
            content = path.read_text(encoding="utf-8", errors="ignore")
        sections.append(f"===== {path.name} =====\n{content}")

    return "\n\n".join(sections), [p.name for p in files]
